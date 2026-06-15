"""Export the LegalEval grading + prompts/responses to readable .docx files.

Usage:
    python -m tools.export_docx runs/<run_id> [--judge claude-opus-4-8] [--out ~/Desktop]

Produces (in --out):
    LegalEval_v7_Grading.docx              — overall table + per-prompt dimension tables + judge notes
    LegalEval_v7_Prompts_and_Responses.docx — per prompt: the prompt text, then each model's full response
"""

import argparse
import csv
import json
import os
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MODELS = ["claude-sonnet-4-6", "gpt-5.5", "gemini-3.5-flash", "qwen-3.7-plus", "llama-4-maverick"]
LABEL = {"claude-sonnet-4-6": "Claude Sonnet 4.6", "gpt-5.5": "GPT-5.5",
         "gemini-3.5-flash": "Gemini 3.5 Flash", "qwen-3.7-plus": "Qwen 3.7 Plus",
         "llama-4-maverick": "Llama 4 Maverick"}
DIMS = ["faithfulness", "directionality", "coverage", "soundness", "actionability"]


def load(run_dir: Path, judge: str):
    suite = json.loads((ROOT / "suite" / "prompts.json").read_text())
    pmeta = {p["id"]: p for p in suite["prompts"]}
    order = [p["id"] for p in suite["prompts"]]
    resp = defaultdict(dict)
    for f in (run_dir / "responses").glob("*.json"):
        r = json.loads(f.read_text())
        resp[r["prompt_id"]][r["provider"]] = r
    grades = defaultdict(dict)
    gdir = run_dir / "grades" / judge
    for f in gdir.glob("*.json"):
        d = json.loads(f.read_text())
        grades[d["prompt_id"]][d["provider"]] = d
    cells = {}
    csv_path = ROOT / "results" / f"{run_dir.name}_{judge}.csv"
    if csv_path.exists():
        for row in csv.DictReader(csv_path.open()):
            cells[(row["prompt_id"], row["provider"])] = row
    return suite, pmeta, order, resp, grades, cells


def _md_paragraphs(doc, text):
    """Render response markdown lightly: #/##/### -> bold headings, keep the rest as paragraphs."""
    for block in text.split("\n"):
        line = block.rstrip()
        if not line:
            continue
        if line.startswith("#"):
            lvl = len(line) - len(line.lstrip("#"))
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip("# ").strip())
            run.bold = True
            run.font.size = Pt(13 - min(lvl, 3))
        elif line.startswith("|"):
            p = doc.add_paragraph(line)
            p.style = doc.styles["No Spacing"]
            for r in p.runs:
                r.font.name = "Courier New"
                r.font.size = Pt(8)
        else:
            doc.add_paragraph(line)


def grading_docx(out, suite, pmeta, order, grades, cells):
    doc = Document()
    doc.add_heading("LegalEval v7 — Grading Report", 0)
    doc.add_paragraph(f"Run {suite.get('suite','v7')} · judge Opus 4.8 · k=1 · rubric {suite['rubric_version']}. "
                      "Scores normalized 0–1 after the 0.40 gate cap. Generation incognito.")
    # overall
    by = defaultdict(list); gt = defaultdict(int); n = defaultdict(int)
    for (pid, prov), r in cells.items():
        by[prov].append(float(r["norm"])); n[prov] += 1
        if r["gates_tripped"]:
            gt[prov] += 1
    doc.add_heading("Overall ranking", 1)
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Model", "Mean score", "Gate-trip rate", "Trips"]):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for m in sorted([m for m in MODELS if n[m]], key=lambda m: -sum(by[m]) / n[m]):
        c = t.add_row().cells
        c[0].text = LABEL[m]; c[1].text = f"{sum(by[m])/n[m]:.3f}"
        c[2].text = f"{gt[m]/n[m]*100:.0f}%"; c[3].text = f"{gt[m]}/{n[m]}"
    # per prompt
    doc.add_heading("Per-prompt grading", 1)
    for pid in order:
        m = pmeta[pid]
        doc.add_heading(f"{pid} · {m.get('name','')}  (/{m.get('max_points','?')}, gates: {', '.join(m.get('gates',[]))})", 2)
        t = doc.add_table(rows=1, cols=9); t.style = "Light Grid Accent 1"
        for i, h in enumerate(["Model", "Faith", "Dir", "Cov", "Sound", "Act", "Raw/Max", "Norm", "Gate cap"]):
            t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for prov in MODELS:
            c = cells.get((pid, prov)); a = grades.get(pid, {}).get(prov, {}).get("aggregate", {})
            row = t.add_row().cells
            row[0].text = LABEL[prov]
            for j, d in enumerate(DIMS):
                v = a.get(d); row[1 + j].text = ("%g" % v) if isinstance(v, (int, float)) else "–"
            if c:
                row[6].text = f"{c['raw']}/{c['max']}"; row[7].text = f"{float(c['norm']):.3f}"
                row[8].text = ("⚠ " + c["gates_tripped"].replace("gate_", "")) if c["gates_tripped"] else "—"
        doc.add_paragraph()
        pn = doc.add_paragraph(); pn.add_run("Judge notes").bold = True
        for prov in MODELS:
            a = grades.get(pid, {}).get(prov, {}).get("aggregate", {})
            note = (a.get("notes") or [""])[0]
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{LABEL[prov]}: ").bold = True
            p.add_run(note)
    doc.save(out)
    return out


def responses_docx(out, suite, pmeta, order, resp):
    doc = Document()
    doc.add_heading("LegalEval v7 — Prompts & Model Responses", 0)
    doc.add_paragraph("Generation only (ungraded) · n=1 · incognito (bare prompt, no system/tools/memory). "
                      "Models: Claude Sonnet 4.6 · GPT-5.5 · Gemini 3.5 Flash · Qwen 3.7 Plus · Llama 4 Maverick.")
    for pid in order:
        m = pmeta[pid]
        doc.add_heading(f"{pid} · {m.get('name','')}  (/{m.get('max_points','?')}, gates: {', '.join(m.get('gates',[]))})", 1)
        ph = doc.add_paragraph(); ph.add_run("▶ PROMPT (sent to every model)").bold = True
        box = doc.add_paragraph(m.get("prompt_text") or "[no prompt_text]")
        box.style = doc.styles["Intense Quote"]
        for prov in MODELS:
            r = resp.get(pid, {}).get(prov)
            doc.add_heading(LABEL[prov], 2)
            if not r:
                doc.add_paragraph("[no response]"); continue
            _md_paragraphs(doc, r["response_text"] or "[empty]")
        doc.add_page_break()
    doc.save(out)
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("run_dir", type=Path)
    ap.add_argument("--judge", default="claude-opus-4-8")
    ap.add_argument("--out", default="~/Desktop")
    args = ap.parse_args()
    out = Path(os.path.expanduser(args.out)); out.mkdir(parents=True, exist_ok=True)
    suite, pmeta, order, resp, grades, cells = load(args.run_dir, args.judge)
    g = grading_docx(out / "LegalEval_v7_Grading.docx", suite, pmeta, order, grades, cells)
    r = responses_docx(out / "LegalEval_v7_Prompts_and_Responses.docx", suite, pmeta, order, resp)
    print("wrote:", g); print("wrote:", r)


if __name__ == "__main__":
    main()
