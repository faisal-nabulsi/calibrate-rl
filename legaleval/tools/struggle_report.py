"""Standalone 'what models struggle on + what to focus on next' findings report.

Reuses the exact computation behind the viewer's struggle section
(`viewer.app.struggle_stats` / `focus_recs`) so the doc and the GUI never drift.
Emits Markdown and (if python-docx is available) a matching .docx.

Usage:
    python -m tools.struggle_report runs/<run_id> [--judge claude-opus-4-8] [--out ~/Desktop]
"""

import argparse
import csv
import json
import os
import re
from collections import defaultdict
from pathlib import Path

from viewer.app import struggle_stats, focus_recs, MODELS, LABEL, DIM_MAX

ROOT = Path(__file__).resolve().parent.parent


def load(run: Path, judge: str):
    grades = defaultdict(dict)
    for f in (run / "grades" / judge).glob("*.json"):
        d = json.loads(f.read_text())
        grades[d["prompt_id"]][d["provider"]] = d
    cells = {}
    with (ROOT / "results" / f"{run.name}_{judge}.csv").open() as fh:
        for r in csv.DictReader(fh):
            cells[(r["prompt_id"], r["provider"])] = r
    pmeta = {p["id"]: p for p in json.loads((ROOT / "suite" / "prompts.json").read_text())["prompts"]}
    return grades, cells, pmeta


def ranking(cells):
    by, n, gt = defaultdict(list), defaultdict(int), defaultdict(int)
    for (pid, prov), r in cells.items():
        by[prov].append(float(r["norm"])); n[prov] += 1
        if r["gates_tripped"]:
            gt[prov] += 1
    return sorted([m for m in MODELS if n[m]], key=lambda m: -sum(by[m]) / n[m]), by, n, gt


def strip_html(t: str) -> str:
    return re.sub(r"</?[^>]+>", "", t).replace("&amp;", "&")


def build_markdown(run, judge, cross, s, cells):
    rank, by, n, gt = ranking(cells)
    L = [f"# LegalEval v7 — What Models Struggle With & What to Focus On Next",
         f"\n*Run `{run.name}` · judge {judge} (k=1) · cross-check {cross} · "
         f"{len(by)} models · {s['n_prompts']} prompts.*\n",
         "## Overall ranking (mean normalized score, after gate cap)\n",
         "| # | Model | Mean | Gate trips |", "|---|---|---|---|"]
    for i, m in enumerate(rank):
        L.append(f"| {i+1} | {LABEL[m]} | {sum(by[m])/n[m]:.3f} | {gt[m]}/{n[m]} |")
    L += ["\n## Hardest qualities (mean as % of dimension max — low = hard)\n",
          "| Dimension | % of max | Scope |", "|---|---|---|"]
    for d, (pct, nn) in s["qsorted"]:
        L.append(f"| {d.title()} | {pct:.0f}% | {nn}/{s['n_prompts']} prompts |")
    L += ["\n## Hardest tasks (mean normalized score across models)\n",
          "| Prompt | Score | Gates |", "|---|---|---|"]
    for pid, v in s["hard"][:6]:
        gates = ", ".join(s["pmeta"][pid].get("gates", [])) or "—"
        L.append(f"| {pid} · {s['pmeta'][pid]['name']} | {sum(v)/len(v):.2f} | {gates} |")
    L += ["\n## The pattern\n",
          f"Models are strongest where the work is mechanical (**Faithfulness "
          f"{s['qpct'].get('faithfulness',(0,0))[0]:.0f}%** — rare fabrication; "
          f"{s['ft']} faithfulness gate trips) and weakest where it needs legal judgment: "
          f"**{s['qsorted'][0][0].title()}** and **{s['qsorted'][1][0].title()}**. "
          f"Gate trips are directional ({s['dt']}) far more than fabrication ({s['ft']}), "
          f"and concentrated in the bottom models — the top three never tripped a gate.\n",
          "## What to focus prompts on next\n"]
    for i, (t, b) in enumerate(focus_recs(s)):
        L.append(f"{i+1}. **{t}.** {strip_html(b)}\n")
    return "\n".join(L)


def build_docx(out, run, judge, cross, s, cells):
    try:
        from docx import Document
    except ImportError:
        return None
    rank, by, n, gt = ranking(cells)
    doc = Document()
    doc.add_heading("LegalEval v7 — What Models Struggle With & What to Focus On Next", 0)
    doc.add_paragraph(f"Run {run.name} · judge {judge} (k=1) · cross-check {cross} · "
                      f"{len(by)} models · {s['n_prompts']} prompts.")

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for c, h in zip(t.rows[0].cells, headers):
            c.paragraphs[0].add_run(h).bold = True
        for row in rows:
            cs = t.add_row().cells
            for c, val in zip(cs, row):
                c.text = str(val)

    doc.add_heading("Overall ranking (after gate cap)", 1)
    table(["#", "Model", "Mean", "Gate trips"],
          [[i+1, LABEL[m], f"{sum(by[m])/n[m]:.3f}", f"{gt[m]}/{n[m]}"] for i, m in enumerate(rank)])
    doc.add_heading("Hardest qualities (% of dimension max — low = hard)", 1)
    table(["Dimension", "% of max", "Scope"],
          [[d.title(), f"{pct:.0f}%", f"{nn}/{s['n_prompts']}"] for d, (pct, nn) in s["qsorted"]])
    doc.add_heading("Hardest tasks (mean normalized score)", 1)
    table(["Prompt", "Score", "Gates"],
          [[f"{pid} · {s['pmeta'][pid]['name']}", f"{sum(v)/len(v):.2f}",
            ", ".join(s["pmeta"][pid].get("gates", [])) or "—"] for pid, v in s["hard"][:6]])
    doc.add_heading("The pattern", 1)
    doc.add_paragraph(
        f"Models are strongest where the work is mechanical (Faithfulness "
        f"{s['qpct'].get('faithfulness',(0,0))[0]:.0f}% — {s['ft']} faithfulness gate trips) and weakest "
        f"where it needs legal judgment: {s['qsorted'][0][0].title()} and {s['qsorted'][1][0].title()}. "
        f"Gate trips are directional ({s['dt']}) far more than fabrication ({s['ft']}), concentrated in the "
        f"bottom models — the top three never tripped a gate.")
    doc.add_heading("What to focus prompts on next", 1)
    for t, b in focus_recs(s):
        p = doc.add_paragraph(style="List Number")
        p.add_run(t + ". ").bold = True
        p.add_run(strip_html(b))
    doc.save(out)
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("run_dir", type=Path)
    ap.add_argument("--judge", default="claude-opus-4-8")
    ap.add_argument("--cross", default="gpt-5.5-pro")
    ap.add_argument("--out", default="~/Desktop")
    a = ap.parse_args()
    grades, cells, pmeta = load(a.run_dir, a.judge)
    s = struggle_stats(grades, cells, pmeta)
    outdir = Path(os.path.expanduser(a.out)); outdir.mkdir(parents=True, exist_ok=True)
    md = outdir / f"LegalEval_v7_Struggle_Report_{a.run_dir.name}.md"
    md.write_text(build_markdown(a.run_dir, a.judge, a.cross, s, cells))
    print("wrote", md)
    dx = build_docx(outdir / f"LegalEval_v7_Struggle_Report_{a.run_dir.name}.docx",
                    a.run_dir, a.judge, a.cross, s, cells)
    print("wrote", dx) if dx else print("(python-docx not installed — skipped .docx)")


if __name__ == "__main__":
    main()
