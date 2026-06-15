"""Build an attorney review packet (.docx) of the cells most worth a lawyer's time.

Flagged = the human-confirm-queue items (borderline gate calls) UNION any cross-judge
gate disagreements (where a second judge's grade dir differs from the primary). Grouped
by prompt; each shows the prompt, the gold key, the flagged model responses + the AI
grade, and blank "Attorney verdict / reasoning" fields. Ends with design questions.

Usage:
    python -m tools.make_review_packet runs/<run_id> [--judge claude-opus-4-8]
        [--cross gpt-5.5] [--out ~/Desktop]
"""

import argparse
import csv
import json
import os
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
LABEL = {"claude-sonnet-4-6": "Claude Sonnet 4.6", "gpt-5.5": "GPT-5.5",
         "gemini-3.5-flash": "Gemini 3.5 Flash", "qwen-3.7-plus": "Qwen 3.7 Plus",
         "llama-4-maverick": "Llama 4 Maverick"}
DIMS = ["faithfulness", "directionality", "coverage", "soundness", "actionability"]
ORDER = ["A1", "A2", "A3", "B1", "B2", "B3", "C1", "C2", "C3", "D1", "D2", "D3", "E1", "E2"]


def flagged_cells(run, judge, cross):
    """Return {(pid,prov): reason}."""
    flags = {}
    cq = ROOT / "results" / f"{run.name}_confirm_queue.json"
    if cq.exists():
        for it in json.loads(cq.read_text()):
            fn = it["file"].split("/")[-1].replace("_s0.json", "")
            pid, prov = fn.split("_", 1)
            g = ",".join(x.replace("gate_", "") for x in it.get("gates_tripped") or [])
            flags[(pid, prov)] = f"flagged for human-confirm ({g + ' gate trip' if g else 'borderline'})"
    # cross-judge directionality disagreements
    cdir = run / "grades" / cross
    if cdir.exists():
        pmeta = {p["id"]: p for p in json.loads((ROOT / "suite" / "prompts.json").read_text())["prompts"]}
        prim = {}
        for r in csv.DictReader((ROOT / "results" / f"{run.name}_{judge}.csv").open()):
            prim[(r["prompt_id"], r["provider"])] = "directionality" in (r["gates_tripped"] or "")
        for f in cdir.glob("*.json"):
            d = json.loads(f.read_text()); pid, prov = d["prompt_id"], d["provider"]
            if "directionality" not in pmeta.get(pid, {}).get("gates", []):
                continue
            cd = d["aggregate"].get("gate_directionality", {}).get("tripped")
            pd = prim.get((pid, prov), False)
            if cd != pd:
                flags[(pid, prov)] = flags.get((pid, prov), "") + \
                    f" | cross-judge SPLIT: {judge}={'trip' if pd else 'pass'}, {cross}={'trip' if cd else 'pass'}"
    return flags


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("run_dir", type=Path)
    ap.add_argument("--judge", default="claude-opus-4-8")
    ap.add_argument("--cross", default="gpt-5.5")
    ap.add_argument("--out", default="~/Desktop")
    a = ap.parse_args()
    run = a.run_dir
    suite = json.loads((ROOT / "suite" / "prompts.json").read_text())
    pmeta = {p["id"]: p for p in suite["prompts"]}
    resp = {}
    for f in (run / "responses").glob("*.json"):
        r = json.loads(f.read_text()); resp[(r["prompt_id"], r["provider"])] = r
    grades = {}
    for f in (run / "grades" / a.judge).glob("*.json"):
        d = json.loads(f.read_text()); grades[(d["prompt_id"], d["provider"])] = d
    cells = {}
    cp = ROOT / "results" / f"{run.name}_{a.judge}.csv"
    for row in csv.DictReader(cp.open()):
        cells[(row["prompt_id"], row["provider"])] = row

    flags = flagged_cells(run, a.judge, a.cross)
    by_prompt = defaultdict(list)
    for (pid, prov), reason in flags.items():
        by_prompt[pid].append((prov, reason))

    doc = Document()
    doc.add_heading("LegalEval v7 — Attorney Review Packet", 0)
    doc.add_paragraph(f"Judge: {a.judge} (k=1) · cross-check: {a.cross} · run {run.name}. "
                      f"{sum(len(v) for v in by_prompt.values())} flagged response(s) across "
                      f"{len(by_prompt)} prompts.")
    # instructions / context
    doc.add_heading("Read this first", 1)
    for t in [
        "GRADE THE INSTANCE, NOT THE WORLD. Score against the contract text shown here only — "
        "not the real underlying deal. If a model is 'right' about the real ACORD/Kubient contract "
        "but wrong about the text below, the text governs.",
        "Gates & cap. Two gates: Faithfulness (fabricating clauses/figures or asserting facts not in "
        "the instance) and Directionality (advocating the wrong party / adding a non-requested term). "
        "ANY gate trip caps the response at 0.40 regardless of quality.",
        "Five dimensions: Faithfulness, Directionality, Coverage, Soundness, Actionability.",
        "WHY YOU: these gold keys were drafted with AI assistance. Your independent legal judgment is the "
        "missing validation. Every place you'd disagree with the AI's 'right answer' is either a gold-key "
        "bug or a defensible position the rubric should credit — that disagreement list is the deliverable.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    for pid in [p for p in ORDER if p in by_prompt]:
        m = pmeta[pid]
        doc.add_page_break()
        doc.add_heading(f"{pid} · {m.get('name','')}  (/{m.get('max_points','?')}, gates: {', '.join(m.get('gates',[]))})", 1)
        doc.add_heading("Prompt (what every model saw)", 3)
        doc.add_paragraph(m.get("prompt_text") or "[none]", style="Intense Quote")
        rp = ROOT / "suite" / m.get("rubric_file", f"rubrics/{pid}.json")
        gold = json.loads(rp.read_text()).get("rubric_markdown", "") if rp.exists() else ""
        doc.add_heading("Gold key / rubric", 3)
        doc.add_paragraph(gold or "[none]", style="Quote")
        for prov, reason in sorted(by_prompt[pid]):
            r = resp.get((pid, prov)); c = cells.get((pid, prov)); g = grades.get((pid, prov), {})
            agg = g.get("aggregate", {})
            doc.add_heading(f"▶ {LABEL.get(prov, prov)}", 2)
            why = doc.add_paragraph(); why.add_run("Why flagged: ").bold = True; why.add_run(reason)
            sc = doc.add_paragraph()
            sc.add_run("AI grade: ").bold = True
            sc.add_run(f"score {float(c['norm']):.2f} ({c['raw']}/{c['max']})"
                       f"{'  ⚠ ' + c['gates_tripped'] if c and c['gates_tripped'] else ''}  |  "
                       + " · ".join(f"{d[:1].upper()}:{agg.get(d,'–')}" for d in DIMS))
            note = (agg.get("notes") or [""])[0]
            if note:
                jn = doc.add_paragraph(); jn.add_run("Judge note: ").bold = True; jn.add_run(note)
            doc.add_paragraph("Model response:").runs[0].bold = True
            doc.add_paragraph(r["response_text"] if r else "[no response]")
            # blank attorney fields
            av = doc.add_paragraph()
            av.add_run("ATTORNEY VERDICT").bold = True
            av.add_run("  — does the AI grade/gate call hold?   [  AGREE  /  DISAGREE  /  PARTLY  ]")
            doc.add_paragraph("Reasoning: ______________________________________________________________")
            doc.add_paragraph("_________________________________________________________________________")
            doc.add_paragraph("If you'd grade differently — your score & why: ___________________________")

    # design questions appendix
    doc.add_page_break()
    doc.add_heading("Design-level questions (the framework itself)", 1)
    for q in [
        "For each gold key here, is the 'correct' answer legally right — and are there reasonable "
        "alternative positions we're penalizing as wrong?",
        "Is Faithfulness rightly the hardest gate, and is Directionality (advocating the wrong party) "
        "appropriately a gate at all?",
        "Is capping a response to 0.40 the right severity for a single wrong-party recommendation?",
        "Do the 14 tasks represent the legal work where AI error actually creates risk? What's missing?",
        "Are the 'modified-from-real' contracts internally consistent, or do edits create ambiguities "
        "that make a graded 'concern' unfair?",
        "Would you send the top-scored drafts (D1 NDA, D2 breach letter) as-is? What would you fix?",
        "Did every model AND our gold key miss something you'd have caught? (gold-key blind spots)",
        "A1 specifically: is the absence of standstill / return-destruction / injunctive-relief / "
        "no-reliance / rep-liability really NEUTRAL-TO-FAVORABLE for the Receiving Party?",
        "E2 specifically: are all six conflicts really bridgeable, and is a stated 2-day termination a "
        "bargaining position rather than a genuine constraint?",
    ]:
        doc.add_paragraph(q, style="List Number")

    outdir = Path(os.path.expanduser(a.out)); outdir.mkdir(parents=True, exist_ok=True)
    out = outdir / "LegalEval_v7_Attorney_Review_Packet.docx"
    doc.save(out)
    print("wrote", out, f"({sum(len(v) for v in by_prompt.values())} responses, {len(by_prompt)} prompts)")


if __name__ == "__main__":
    main()
