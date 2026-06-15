"""Light Markdown -> nicely-formatted .docx converter (headings, bullets, numbered
lists, **bold**, and pipe-tables rendered as real Word tables).

Usage:
    python -m tools.md_to_docx INPUT.md OUTPUT.docx ["Optional Title"]
"""

import re
import sys

from docx import Document
from docx.shared import Pt


def _add_runs(paragraph, text):
    """Render **bold** inline; leave the rest plain."""
    for i, chunk in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        else:
            paragraph.add_run(chunk)


def _table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]  # row 1 is the |---| separator
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h.replace("**", ""))
        run.bold = True
    for r in body:
        rc = t.add_row().cells
        for i in range(len(header)):
            _add_runs(rc[i].paragraphs[0], r[i] if i < len(r) else "")
    doc.add_paragraph()


def convert(md_text, out_path, title=None):
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # table block
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"\s*\|[-:\s|]+\|\s*$", lines[i + 1]):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i]); i += 1
            _table(doc, block)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).replace("**", ""), min(len(m.group(1)), 4))
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet"); _add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+[.)]\s+", line):
            p = doc.add_paragraph(style="List Number"); _add_runs(p, re.sub(r"^\s*\d+[.)]\s+", "", line))
        elif re.match(r"^>\s?", line):
            p = doc.add_paragraph(style="Quote"); _add_runs(p, re.sub(r"^>\s?", "", line))
        else:
            p = doc.add_paragraph(); _add_runs(p, line)
        i += 1
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src, dst = sys.argv[1], sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else None
    print("wrote", convert(open(src).read(), dst, title))
