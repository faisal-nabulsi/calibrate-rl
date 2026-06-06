"""
Generate the mentor-review proposal docx for the depth-0 goldilocks run.
Writes to the OUTER docs folder. python3 tools/make_mentor_doc.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
OUTER = os.path.dirname(os.path.dirname(HERE))
OUT = os.path.join(OUTER, "CalibrateRL_v10_depth0_training_proposal.docx")
GREY = RGBColor(0x6b, 0x6b, 0x76)
ACC = RGBColor(0x32, 0x57, 0xd6)


def main():
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(10.5)

    d.add_heading("CalibrateRL — Depth-0 Goldilocks Training Run", 0)
    s = d.add_paragraph()
    s.add_run("Proposal for review before launch · Qwen2.5-7B · GRPO/LoRA").italic = True

    def h(t): d.add_heading(t, level=1)
    def p(t): return d.add_paragraph(t)
    def b(items):
        for it in items:
            d.add_paragraph(it, style="List Bullet")
    def note(t):
        par = d.add_paragraph(); r = par.add_run(t); r.font.color.rgb = GREY; return par

    # ── 1. TL;DR ────────────────────────────────────────────────────────────
    h("1. The ask")
    p("We're proposing a small GRPO hillclimb to validate the core CalibrateRL "
      "loop end-to-end: calibrate a model's difficulty → keep only the goldilocks-"
      "zone problems (2–6/8 for the 7B) → train on those → watch a held-out "
      "goldilocks slice improve. This run is intentionally small and is about "
      "proving the loop works and produces signal, not about a headline score. "
      "We'd value your read on the design before we launch.")
    note("Premise already validated in calibration: capability ordering holds — "
         "1.5B mean reward 0.165 vs 7B 0.508 (3.1×) on the same set — so the harness "
         "measures real skill. The 1.5B was a calibration tier only (not trained); "
         "this run trains the 7B; a 3B with its own calibrated curriculum is a later, "
         "separate effort.")

    # ── 2. The run ──────────────────────────────────────────────────────────
    h("2. The run (proposed config)")
    b([
        "Model: Qwen2.5-7B-Instruct, LoRA (r=32, α=64, dropout 0.05, all-linear).",
        "Algorithm: GRPO (TRL), 8 generations/prompt, temperature 1.0 (constant — "
        "matched to the calibration sampling temperature; no annealing), "
        "max completion 1024, KL β=0.1, DAPO loss, truncated-completion masking.",
        "Data: 106 goldilocks-zone depth-0 skeleton problems (2–6/8 for the 7B), "
        "spanning all 26 concepts that have goldilocks instances. Train set only — "
        "the 12 held-out goldilocks are never trained on.",
        "Effective batch: per_device 2 × grad-accum 16 = 32 completions/step = "
        "4 unique prompts/step. Because the set is all-goldilocks, ~0% of prompts "
        "are zero-gradient (vs ~38% if we'd trained the full unfiltered set).",
        "Schedule: max_steps 120 ≈ 4.5 epochs over 106 at 4 prompts/step. "
        "Checkpoint + held-out eval every 27 steps (≈ once per epoch).",
        "Reward: binary correctness (1/0) + a small format term (+0.1 for boxed "
        "answer with work, −0.2 for no work), via the hardened grader (below).",
    ])
    note("Success criterion: held-out goldilocks pass@8 rises over the run without "
         "the format/length tripwires degrading. That's the loop producing signal.")

    # ── 3. Held-out design ──────────────────────────────────────────────────
    h("3. Held-out & monitoring design")
    b([
        "Per-step monitor = 12 held-out goldilocks (1 per concept, from the 12 best-"
        "populated concepts; stratified, seed 42; disjoint from train, overlap 0). "
        "We can't stratify a ~10% held-out across all 26 concepts — most have too "
        "few goldilocks to spare one — so we cover the well-populated ones and keep "
        "the scarce concepts entirely in train.",
        "Metric = pass@8 @ temp 1.0 (not greedy). On 12 problems greedy is coarse, "
        "and these are goldilocks (model is ~half-right by construction), so pass@8 "
        "directly measures the pass-rate we expect training to lift. Cost: 96 "
        "generations/eval — trivial. Evaluated at step 0 (baseline), every epoch, and end.",
        "Reward-hacking tripwires logged each eval: boxed-rate (format drift) and "
        "mean completion length (length gaming / degeneration).",
        "AMC transfer eval is intentionally OFF for this run. It's the real-target "
        "capability claim, so we don't want to steer a tiny hillclimb on it (test-set "
        "peeking). If we want the AMC transfer curve later, we rebuild it post-hoc "
        "from the saved per-epoch checkpoints — no mid-run decision touches AMC.",
    ])

    # ── 4. What changed to get here ─────────────────────────────────────────
    h("4. What we changed to make this run honest")
    b([
        "Grader: hardened and consolidated to ONE tested source (reward_func.py). "
        "Now: takes the LAST boxed answer (credits self-correction), parses "
        "\\frac/\\dfrac, and gates the bare-number fallback so candidate-listing "
        "(\"could be 10, 20, or 42\") earns no reward — GRPO would otherwise learn to "
        "exploit it. 21-case regression test. The 4 stale eval-script copies were "
        "deleted/repointed so eval and training share one rubric.",
        "Goldilocks selection: built by RE-GRADING the stored calibration rollouts "
        "with the current hardened grader (not by trusting the old stored zones). "
        "This shifted the goldilocks set by 2 (120→118) — 6 problems re-graded "
        "differently, 2 crossed out of the 2–6/8 band — so selection is now "
        "consistent with what training actually rewards.",
        "Config correction: train_grpo previously pointed at the full unfiltered set "
        "(3,849 problems) at 500 steps — which is ~0.5 epoch with ~38% zero-gradient "
        "prompts, and not the calibration thesis. Now it trains the goldilocks subset "
        "at ~4.5 epochs. (We also fixed an earlier temp/dataset mismatch: it had been "
        "set to train on an old dataset at an annealing temperature ≠ calibration.)",
        "Human-eval UI: a local 2-page app — (1) call the model live on any held-out "
        "problem, auto-graded by the real grader; (2) browse the 12 held-out with "
        "their 8 rollouts, pass-rate, and zone. Plus a .docx transcript of the base "
        "model on the 12 (attached separately) for the team to read.",
    ])

    # ── 5. Rigor note: data integrity ───────────────────────────────────────
    h("5. Data-integrity check (rigor note — 0 impact on this run)")
    p("We built an independent gold-answer checker (recomputes answers from problem "
      "text for ~22 of 29 concepts) and found a real generator bug: the "
      "continued_fraction skeleton rendered problems at depth D but computed answers "
      "at depth D+2 (verified — every wrong gold equals the depth+2 value). It affects "
      "14 rows of the full dataset. None of those rows were in the random-300 "
      "calibration sample, so this run's goldilocks set is unaffected — but it de-risks "
      "scaling, and the checker now gates any future dataset.")

    # ── 6. Caveat ───────────────────────────────────────────────────────────
    h("6. Caveat: OpenRouter baseline ≠ local training weights")
    p("The human-eval baseline (UI + transcript) is served via OpenRouter "
      "(qwen/qwen-2.5-7b-instruct): ~51% overall on the 12 held-out, which re-sample "
      "as goldilocks 9 / borderline 2 / too-easy 1 / too-hard 0 — confirming they sit "
      "in the band. But OpenRouter's serving is not byte-identical to the local "
      "HuggingFace weights we train (local calibration measured ~40% goldilocks on a "
      "different sample). So: use the OpenRouter baseline only for reading reasoning "
      "quality with consistent serving (base-OpenRouter vs trained-checkpoint-OpenRouter); "
      "the actual training delta comes from the in-run local pass@8 monitor. We won't "
      "compare the two pass-rates directly.")

    # ── 7. Deferred ─────────────────────────────────────────────────────────
    h("7. Deliberately deferred (not part of this run)")
    b([
        "AMC transfer measurement (rebuildable post-hoc from checkpoints).",
        "A larger goldilocks pool via calibrating more than 300 problems (120 is "
        "what we've labeled, not the dataset's total goldilocks).",
        "Full-dataset dedup + gold cleanup for future, larger runs.",
        "Effective-batch / step-budget tuning beyond this first hillclimb.",
        "A separate 3B run with its own calibrated curriculum.",
    ])

    # ── 8. Where we'd value your read ───────────────────────────────────────
    h("8. Two things we'd value your read on")
    b([
        "Step budget: ~4.5 epochs over 106 goldilocks — reasonable, or do you expect "
        "overfitting before the held-out moves? Would you rather fewer epochs + a "
        "bigger pool (more calibration first)?",
        "Effective batch: 4 unique prompts/step (32 completions). For all-goldilocks "
        "GRPO is that steady enough, or would you bump grad-accum for less noisy "
        "updates at the cost of wall-clock?",
    ])

    d.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
