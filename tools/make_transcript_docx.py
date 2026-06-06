"""
Build a readable .docx transcript of the goldilocks held-out rollouts for the
team. Mostly the bare transcript (all 8 rollouts per problem), with a light
auto-generated analysis line each (pass-rate + answer spread + one observation).

Reads tools/holdout_results.json (from run_holdout.py); writes the .docx to the
OUTER folder (docs live outside the repo).

    python3 tools/make_transcript_docx.py
"""
import os, json
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
RES = os.path.join(HERE, "holdout_results.json")
OUTER = os.path.dirname(os.path.dirname(HERE))          # tools/ -> repo -> outer folder
OUT = os.path.join(OUTER, "CalibrateRL_v10_goldilocks_holdout_transcript.docx")

GREEN = RGBColor(0x13, 0x7a, 0x37)
RED = RGBColor(0xb3, 0x26, 0x1e)
GREY = RGBColor(0x6b, 0x6b, 0x76)


def light_analysis(r):
    rolls = r["rollouts"]
    preds = Counter(str(x["pred"]) for x in rolls)
    dist = ", ".join(f"{p}×{c}" for p, c in preds.most_common())
    nc = r["n_correct"]
    wrong_preds = Counter(str(x["pred"]) for x in rolls if not x["correct"])
    if nc == 0:
        obs = "Never reached the gold."
    elif nc == r["n"]:
        obs = "Solved every rollout."
    elif len(wrong_preds) == 1:
        obs = (f"The failures all land on a single wrong answer "
               f"({list(wrong_preds)[0]}) — a systematic slip, not noise.")
    elif len(wrong_preds) >= max(3, nc):
        obs = (f"Failures scattered across {len(wrong_preds)} different answers — "
               f"execution/arithmetic noise rather than a method error.")
    else:
        obs = f"Mixed: {len(wrong_preds)} distinct wrong answers among the misses."
    return f"{nc}/{r['n']} correct ({r.get('zone','')}). Final answers: {dist}. {obs}"


def main():
    data = json.load(open(RES))
    meta, results = data["meta"], data["results"]
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    h = doc.add_heading("CalibrateRL · v10 Goldilocks Held-out — Base Model Transcript", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"{meta['model']} · {meta['rollouts']} rollouts/problem · temp {meta['temperature']} "
                f"· {meta['n_problems']} held-out goldilocks problems").italic = True
    z = meta.get("zones", {})
    doc.add_paragraph(
        f"Overall pass-rate {meta['overall_pass_rate']:.0%}.  Re-sampled zones on this set: "
        f"goldilocks {z.get('goldilocks',0)} · borderline {z.get('borderline',0)} · "
        f"too-easy {z.get('too_easy',0)} · too-hard {z.get('too_hard',0)}.")
    note = doc.add_paragraph()
    note.add_run(
        "These are the 12 held-out goldilocks problems the RL run monitors but never "
        "trains on. This is the BASE model (pre-RL) baseline — read it now, then compare "
        "the same set after training to judge whether the reasoning matures. Served via "
        "OpenRouter, which may differ slightly from the local training weights.").font.color.rgb = GREY
    doc.add_paragraph()

    for i, r in enumerate(results, 1):
        hd = doc.add_heading(f"{i}.  {r['concept']}", level=1)
        meta_p = doc.add_paragraph()
        meta_p.add_run(f"gold = {r['gold']}").bold = True
        a = doc.add_paragraph()
        a.add_run("Analysis: ").bold = True
        a.add_run(light_analysis(r))
        doc.add_paragraph(r["problem"]).runs[0].italic = True

        for j, ro in enumerate(r["rollouts"]):
            p = doc.add_paragraph()
            tag = p.add_run(f"— rollout {j}  [{'CORRECT' if ro['correct'] else 'WRONG'}]  "
                            f"(grader read {ro['pred']} via {ro['method']})")
            tag.bold = True
            tag.font.color.rgb = GREEN if ro["correct"] else RED
            doc.add_paragraph(ro["text"])
        doc.add_paragraph()

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"{len(results)} problems, {sum(len(r['rollouts']) for r in results)} rollouts")


if __name__ == "__main__":
    main()
